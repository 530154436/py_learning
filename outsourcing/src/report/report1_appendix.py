# -*- coding: utf-8 -*-
from typing import List

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches, Cm
from pandas import DataFrame
from scipy import stats
from analysis import RESEARCH_TYPE_MAPPING
from config import DATASET_DIR, OUTPUT_DIR
from utils.doc_tpl_util import merge_table_column, set_table_column_font, set_heading_font_style


def appendix_1(doc: Document):
    """
    附表1. 获奖人按研究类型归类
    """
    input_file = DATASET_DIR.joinpath("S2.2-学者关联信息表-对照分组.xlsx")
    df = pd.read_excel(input_file)
    df = df[df["学者类型（获奖人=1，0=对照学者）"] == 1]
    df = df[["姓名", "研究领域", "研究类型（1=基础科学，0=工程技术，2=前沿交叉）"]]
    df['研究类型名称'] = df['研究类型（1=基础科学，0=工程技术，2=前沿交叉）'].map(lambda x: RESEARCH_TYPE_MAPPING[str(x)])
    print(df)

    # 按研究领域分组统计
    df = df.groupby(['研究类型名称', '研究领域']).agg(
        获奖人数量=('姓名', 'count'),
        获奖人名单=('姓名', lambda x: '\n'.join(x))
    ).reset_index()
    print(df)

    # 创建表格（先只建表头）
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'  # 可选样式

    # 设置表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '研究类型'
    hdr_cells[1].text = '研究领域'
    hdr_cells[2].text = '获奖人数量（人）'
    hdr_cells[3].text = '获奖人名单'

    # 存储每个“研究类型”开始的行索引，用于后续合并
    type_start_rows = {}
    current_type = None
    row_index = 1  # 当前是第1行（表头是第0行）

    # 第一步：遍历数据，添加所有行
    for _, row in df.iterrows():
        if row['研究类型名称'] != current_type:
            current_type = row['研究类型名称']
            type_start_rows[current_type] = row_index  # 记录这个类型的第一行索引

        # 添加数据行
        new_row = table.add_row()
        new_row.cells[1].text = row['研究领域']
        new_row.cells[2].text = str(row['获奖人数量'])
        new_row.cells[3].text = row['获奖人名单'].replace('\n', '\n')

        row_index += 1

        # 添加合计行（在每个类型结束后）
        if _ == df.index[-1] or (df['研究类型名称'].iloc[_ + 1] != current_type if _ < len(df) - 1 else True):
            total_row = table.add_row()
            total_row.cells[1].text = '合计'
            total_row.cells[2].text = str(df[df['研究类型名称'] == current_type]['获奖人数量'].sum())
            total_row.cells[3].text = '/'
            row_index += 1

    # 添加数据行：合计
    new_row = table.add_row()
    new_row.cells[0].text = '/'
    new_row.cells[1].text = '总计'
    new_row.cells[2].text = str(df['获奖人数量'].sum())
    new_row.cells[3].text = '/'

    # 第二步：合并“研究类型”列的单元格
    current_type = None
    for i, row in enumerate(table.rows[1:], start=1):  # 跳过表头
        type_text = row.cells[1].text  # 通过“研究领域”判断是否是新类型的第一行

        # 如果当前行是某个研究类型的起始行
        if type_text in df['研究领域'].values and df[df['研究领域'] == type_text]['研究类型名称'].iloc[0] != current_type:
            current_type = df[df['研究领域'] == type_text]['研究类型名称'].iloc[0]
            start_row = i
            # 找到该类型对应的合计行（下个类型开始前）
            next_type_start = None
            for j in range(i + 1, len(table.rows)):
                if table.rows[j].cells[1].text in df['研究领域'].values or table.rows[j].cells[1].text == '合计':
                    next_type_start = j
                    break
            end_row = next_type_start - 1  # 合并到“合计”前一行

            # 合并 cells[0] 从 start_row 到 end_row
            if end_row > start_row:
                cell_to_merge = table.cell(start_row, 0)
                cell_end = table.cell(end_row, 0)
                cell_to_merge.merge(cell_end)
            table.rows[start_row].cells[0].text = current_type
    return table


def appendix_2(doc: Document):
    """
    附表2. 获奖人获奖前后5年论文专利指标数据
    """
    def calculate_stats(df, time_window):
        """ 计算统计值
        """
        stats_df = df[df['时间窗口（0=获奖前5年，1=获奖后5年）'] == time_window]
        stats = {
            '发文量': ['总发文量'],
            '通讯作者论文数(A1)': ['通讯作者论文数（A1）'],
            '专利族数量(A2)': ['专利族数量（A2）'],
            '第一发明人授权专利数量(A3)': ['第一发明人授权专利数量（A3）'],
            '论文篇均被引频次(B1)': ['论文篇均被引频次（B1）'],
            'Q1区通讯作者论文数量占比(B2)': ['前10%高影响力期刊或会议通讯作者论文数量占比（B3）'],
            '单篇最高被引频次(B3)': ['单篇最高被引频次（B2）'],
            '专利被引频次(B4)': ['专利被引频次（B4）']
        }
        result = {}
        for key, cols in stats.items():
            data = stats_df[cols].values.flatten()
            result[key] = {
                '最小值': data.min(),
                '最大值': data.max(),
                '平均值': data.mean(),
                '标准差': data.std()
            }
        return result

    def fill_table(table, stats, time_window_label):
        """ 填充表格数据，并合并第一列的时间窗口单元格
        """
        start_row = len(table.rows)  # 当前最后一行索引（新行将从此开始）

        # 先添加所有数据行
        for key, values in stats.items():
            row = table.add_row()
            cells = row.cells
            cells[1].text = key
            cells[2].text = str(values['最小值'])
            cells[3].text = str(values['最大值'])
            cells[4].text = f"{values['平均值']:.2f}"
            cells[5].text = f"{values['标准差']:.2f}"

        # 计算结束行
        end_row = len(table.rows) - 1  # 最后一行索引

        # 合并第一列中从 start_row 到 end_row 的所有单元格
        if start_row <= end_row:
            cell_start = table.cell(start_row, 0)
            cell_end = table.cell(end_row, 0)
            merged_cell = cell_start.merge(cell_end)
            merged_cell.text = time_window_label

        for i, row in enumerate(table.rows):
            row.cells[0].width = Cm(4)  # 设置第2列（指标列）的宽度为6厘米
            row.cells[1].width = Cm(12)  # 设置第2列（指标列）的宽度为6厘米

    input_file = OUTPUT_DIR.joinpath("A2-评价指标数据集.xlsx")
    _df = pd.read_excel(input_file)
    _df = _df[_df["学者类型（获奖人=1，0=对照学者）"] == 1]

    stats_pre_award = calculate_stats(_df, 0)
    stats_post_award = calculate_stats(_df, 1)

    # 创建表格
    _table = doc.add_table(rows=1, cols=6)
    _table.style = 'Table Grid'  # 可选样式
    hdr_cells = _table.rows[0].cells
    hdr_cells[0].text = '时间窗口'
    hdr_cells[1].text = '指标'
    hdr_cells[2].text = '最小值'
    hdr_cells[3].text = '最大值'
    hdr_cells[4].text = '平均值'
    hdr_cells[5].text = '标准差'

    # 填充数据
    fill_table(_table, stats_pre_award, "获奖前5年")
    fill_table(_table, stats_post_award, "获奖后5年")
    return _table


def appendix_3(doc: Document):
    """
    附表3. 获奖人获奖前后5年论文专利指标数据
    """
    def init_table():
        """ 创建表格和表头
        """
        table = doc.add_table(rows=0, cols=10)  # 先不加行，手动控制
        table.style = 'Table Grid'

        # -----------------------------
        # 第一层表头：时间窗口分组（仅合并“前5年”和“后5年”部分）
        # -----------------------------
        header_row1 = table.add_row().cells
        header_row1[0].text = "序号"
        header_row1[1].text = "获奖人姓名"
        header_row1[2].text = "领域"
        # 第3-6列：获奖前5年差值 → 合并3列
        cell_pre = header_row1[3]
        cell_pre.merge(header_row1[4]).merge(header_row1[5])
        cell_pre.text = "获奖前5年差值"
        for paragraph in cell_pre.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # 第7-10列：获奖后5年差值 → 合并4列
        cell_post = header_row1[6]
        cell_post.merge(header_row1[7]).merge(header_row1[8])
        cell_post.text = "获奖后5年差值"
        for paragraph in cell_post.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        header_row1[9].text = "成长模式"
        # -----------------------------
        # 第二层表头：具体指标名称
        # -----------------------------
        header_row2 = table.add_row().cells
        header_row2[0].merge(header_row1[0])
        header_row2[1].merge(header_row1[1])
        header_row2[2].merge(header_row1[2])
        header_row2[3].text = "综合分数"
        header_row2[4].text = "学术生产力"
        header_row2[5].text = "学术影响力"
        header_row2[6].text = "综合分数"
        header_row2[7].text = "学术生产力"
        header_row2[8].text = "学术影响力"
        header_row2[9].merge(header_row1[9])
        return table

    def fill_table(table, df):
        """ 填充表格数据
        """
        # 只保留获奖人（学者类型 = 1）
        df_awardees = df[df['学者类型（获奖人=1，0=对照学者）'] == 1].reset_index(drop=True)
        for idx, row in df_awardees.iterrows():
            data_row = table.add_row().cells

            data_row[0].text = str(idx + 1)  # 序号
            data_row[1].text = row['姓名']  # 姓名
            data_row[2].text = row['研究领域']  # 领域
            # 获奖前5年差值（直接使用已有字段）
            data_row[3].text = f"{row['差值-综合分数0']:.2f}"  # 综合分数
            data_row[4].text = f"{row['差值-学术生产力0']:.2f}"  # 学术生产力
            data_row[5].text = f"{row['差值-学术影响力0']:.2f}"  # 学术影响力
            # 获奖后5年差值
            data_row[6].text = f"{row['差值-综合分数1']:.2f}"
            data_row[7].text = f"{row['差值-学术生产力1']:.2f}"
            data_row[8].text = f"{row['差值-学术影响力1']:.2f}"
            data_row[9].text = row['成长模式']

        # 设置列宽（可选，提升可读性）
        col_widths = [
            None,  # 序号 → 稍宽一点
            Cm(3.5),  # 姓名
            Cm(4.5),  # 领域
            None,  # 综合分数（前）
            None,  # 学术生产力（前）
            None,  # 学术影响力（前）
            None,  # 综合分数（后）
            None,  # 学术生产力（后）
            None,  # 学术影响力（后）
            Cm(4.5),  # 成长模式 → 更宽
        ]
        for i, width in enumerate(col_widths):
            if not width:
                continue
            for row in table.rows:
                row.cells[i].width = width

    input_file = OUTPUT_DIR.joinpath("A3-差值分析数据集.xlsx")
    _df = pd.read_excel(input_file)
    _table = init_table()
    fill_table(_table, _df)
    return _table


def appendix_4(doc: Document):
    """
    附表3. 获奖人获奖前后5年论文专利指标数据
    """
    def init_table():
        """ 创建表格和表头
        """
        table = doc.add_table(rows=0, cols=9)  # 先不加行，手动控制
        table.style = 'Table Grid'

        # -----------------------------
        # 第一层表头：时间窗口分组（仅合并“前5年”和“后5年”部分）
        # -----------------------------
        header_row1 = table.add_row().cells
        header_row1[0].text = "配对样本：\r获奖后5年-\r获奖前5年"
        header_row1[1].text = "均值差值"
        header_row1[2].text = "标准差"
        header_row1[3].text = "均值的标准误差"
        # 第3-6列：获奖前5年差值 → 合并3列
        cell = header_row1[4]
        cell.merge(header_row1[4]).merge(header_row1[5])
        cell.text = "差分的95%\r置信区间"
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        header_row1[6].text = "t"
        header_row1[7].text = "df"
        header_row1[8].text = "Sig.(双侧)"
        # -----------------------------
        # 第二层表头：具体指标名称
        # -----------------------------
        header_row2 = table.add_row().cells
        header_row2[0].merge(header_row1[0])
        header_row2[1].merge(header_row1[1])
        header_row2[2].merge(header_row1[2])
        header_row2[3].merge(header_row1[3])
        header_row2[4].text = "下限"
        header_row2[5].text = "上限"
        header_row2[6].merge(header_row1[6])
        header_row2[7].merge(header_row1[7])
        header_row2[8].merge(header_row1[8])
        return table

    def calculate_stats(data: DataFrame):
        """ 计算每位获奖人自身的“获奖前后变化” 置信区间和T检验
        """
        # 只保留获奖人（学者类型 = 1）
        df_awardees = data[data['学者类型（获奖人=1，0=对照学者）'] == 1].copy()
        df_awardees['综合分数变化'] = df_awardees['综合分数1'] - df_awardees['综合分数0']
        df_awardees['学术生产力变化'] = df_awardees['学术生产力1'] - df_awardees['学术生产力0']
        df_awardees['学术影响力变化'] = df_awardees['学术影响力1'] - df_awardees['学术影响力0']

        # 提取变化值用于 t 检验
        changes = {
            '综合分数1-\r综合分数0': df_awardees['综合分数变化'],
            '学术生产力1-\r学术生产力0': df_awardees['学术生产力变化'],
            '学术影响力1-\r学术影响力0': df_awardees['学术影响力变化']
        }
        results = []
        for label, change_data in changes.items():
            n = len(change_data)
            mean_diff = change_data.mean()
            std_dev = change_data.std()
            std_err = std_dev / np.sqrt(n)

            # 95% 置信区间
            ci = stats.t.interval(0.95, df=n - 1, loc=mean_diff, scale=std_err)
            ci_lower, ci_upper = ci

            # t 检验：H0: 均值变化 = 0
            t_stat, p_value = stats.ttest_1samp(change_data, 0)

            results.append({
                '指标': label,
                '均值差值': mean_diff,
                '标准差': std_dev,
                '均值的标准误差': std_err,
                'CI下限': ci_lower,
                'CI上限': ci_upper,
                't': t_stat,
                'df': n - 1,
                'Sig.(双侧)': p_value
            })
        return results

    input_file = OUTPUT_DIR.joinpath("A3-差值分析数据集.xlsx")
    df = pd.read_excel(input_file)
    rows = calculate_stats(df)

    # 创建表格
    table = init_table()
    for res in rows:
        row = table.add_row().cells
        row[0].text = res['指标']
        row[1].text = f"{res['均值差值']:.2f}"
        row[2].text = f"{res['标准差']:.2f}"
        row[3].text = f"{res['均值的标准误差']:.2f}"

        # CI 下限和上限填入第4、5列
        row[4].text = f"{res['CI下限']:.2f}"
        row[5].text = f"{res['CI上限']:.2f}"

        row[6].text = f"{res['t']:.2f}"
        row[7].text = str(res['df'])
        row[8].text = f"{res['Sig.(双侧)']:.3f}"

        row[0].width = Cm(5)

    return table


def add_all_appendix_tables():
    doc = Document()

    # 附表
    set_heading_font_style(doc.add_paragraph('附件'))
    set_heading_font_style(doc.add_heading('附表1. 获奖人按研究类型归类'))
    appendix_1(doc)
    set_heading_font_style(doc.add_heading('附表2. 获奖人获奖前后5年论文专利指标数据'))
    appendix_2(doc)
    set_heading_font_style(doc.add_heading('附表3. 获奖人获奖前后5年学术能力指标与对照学者均值的差值'))
    appendix_3(doc)
    set_heading_font_style(doc.add_heading('附表4. 获奖人获奖前后5年学术能力指标配对样本t检验结果'))
    appendix_4(doc)

    # 获取表格并合并单元格
    for i, table in enumerate(doc.tables):
        table = doc.tables[i]
        doc.tables[i] = merge_table_column(table, col_idx=0)  # 合并第 0 列（时间窗口列）
        set_table_column_font(table)

    doc.save('1-首届获奖人获奖前后学术能力量化评估综合报告-附表.docx')
    print("1-首届获奖人获奖前后学术能力量化评估综合报告-附表.docx")


if __name__ == '__main__':
    add_all_appendix_tables()
