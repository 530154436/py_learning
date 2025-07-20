# -*- coding: utf-8 -*-
import pandas as pd
from docx.shared import Mm
from docxtpl import DocxTemplate, InlineImage

from utils.doc_tpl_util import merge_table_column
from utils.plot_util import dataframe_to_chart


def main():
    doc = DocxTemplate("data/demo/WordTemplate.docx")
    context = dict()

    # 1.变量替换
    context['report_name'] = "综合评估报告"

    # 2.表格填充
    metrics = [
        {"c1": "获奖前5年", "c2": "学术生产力", "c3": 0.30, "c4": 32.27, "c5": 11.58, "c6": 8.28},
        {"c2": "学术影响力", "c3": 0, "c4": 40, "c5": 16.19, "c6": 8.42},
        {"c2": "国际引领力", "c3": 0, "c4": 11.88, "c5": 5.32, "c6": 2.68},
        {"c2": "综合分数", "c3": 1.11, "c4": 60.42, "c5": 33.09, "c6": 15.01},
        {"c1": "获奖后5年", "c2": "学术生产力", "c3": 0, "c4": 32.04, "c5": 11.71, "c6": 7.92},
        {"c2": "学术影响力", "c3": 0, "c4": 41.60, "c5": 19.93, "c6": 8.99},
        {"c2": "国际引领力", "c3": 0, "c4": 12.31, "c5": 5.96, "c6": 3.30},
        {"c2": "综合分数", "c3": 3, "c4": 69.26, "c5": 37.60, "c6": 16.42},
    ]
    context.update({"metrics": metrics})

    # 3.导入excel表格数据并绘制折线图嵌入到doc
    df = pd.read_excel("data/demo/data.xlsx")
    image_path = "data/demo/image.png"
    dataframe_to_chart(df, x_col="姓名", y_cols=["综合分数"], title="综合分数", image_path=image_path)
    insert_image = InlineImage(doc, image_path, width=Mm(140))
    context["insert_image"] = insert_image

    # 渲染
    doc.render(context)

    # 获取表格并合并单元格
    for i, table in enumerate(doc.tables):
        table = doc.tables[0]  # 假设要处理的是第一个表格
        doc.tables[i] = merge_table_column(table, col_idx=0)  # 合并第 0 列（时间窗口列）

    # 保存
    doc.save("data/demo/generated_doc.docx")


if __name__ == '__main__':
    main()
